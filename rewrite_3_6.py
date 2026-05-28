import docx
from docx.enum.text import WD_COLOR_INDEX

def update_3_6_final(file_path, output_path):
    doc = docx.Document(file_path)
    
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("3.6 Use Case Analysis"):
            start_idx = i
        elif start_idx != -1 and (p.text.strip().startswith("3.7") or p.text.strip().startswith("Chapter 4")):
            end_idx = i
            break
            
    if start_idx == -1:
        print("Could not find section 3.6")
        return

    # Delete existing paragraphs between 3.6 and 3.7
    target_end = end_idx if end_idx != -1 else len(doc.paragraphs)
    for i in range(target_end - 1, start_idx, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
        doc._body._pelfree_cache = None

    # Re-find the next section to insert before it
    ch3_7 = None
    if end_idx != -1:
        for p in doc.paragraphs:
            if p.text.strip().startswith("3.7") or p.text.strip().startswith("Chapter 4"):
                ch3_7 = p
                break

    use_cases = [
        ("Use Case 1: Asking a Question", 
         "A user wants to know about PlagPro's services. They type a message in the chat box. The AI (Dialogflow) identifies what the user wants and sends back a helpful answer immediately. If the AI is confused, it asks the user to rephrase or contact support."),
        
        ("Use Case 2: Creating a Support Ticket", 
         "When a user needs formal help, they use the 'Create Ticket' feature. A simple form appears inside the chat window where they enter their name and issue. The system saves this in the MongoDB database and gives the user a reference number like 'EMP-1001'."),
        
        ("Use Case 3: Checking Ticket Status", 
         "A user with a ticket number can check its progress at any time. They enter the number in the sidebar's search field. The system looks up the latest information (like 'Resolved' or 'In Progress') and displays it directly to the user."),
        
        ("Use Case 4: Human Handoff (Real-time Help)", 
         "If the AI cannot solve a problem, the user can request a human agent. The system creates a special 'handoff' token. This alerts an administrator on their dashboard, allowing them to take over the conversation and chat with the user in real-time."),
        
        ("Use Case 5: Admin Dashboard Management", 
         "Administrators log into a private dashboard to see all user activity. They can see charts showing how many people are using the bot, read all support tickets, and send manual replies to users who need expert guidance.")
    ]

    # Insert before the next section
    target = ch3_7 if ch3_7 else doc.paragraphs[-1]
    
    # We add them in order by using insert_paragraph_before on the target section
    for title, desc in use_cases:
        # Title
        p_title = target.insert_paragraph_before(title)
        p_title.runs[0].bold = True
        p_title.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # Description
        p_desc = target.insert_paragraph_before(desc)
        for run in p_desc.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        # Spacing
        target.insert_paragraph_before("")

    doc.save(output_path)
    print(f"Section 3.6 simplified and saved to {output_path}")

if __name__ == "__main__":
    update_3_6_final("Project Report (Updated v3).docx", "Project Report (Updated v4).docx")
