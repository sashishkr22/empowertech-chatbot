import docx
from docx.enum.text import WD_COLOR_INDEX

def update_3_6_simple(file_path, output_path):
    doc = docx.Document(file_path)
    
    start_idx = -1
    end_idx = -1
    
    # 1. Find the section boundaries
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("3.6 Use Case Analysis"):
            start_idx = i
        elif start_idx != -1 and (p.text.strip().startswith("3.7") or p.text.strip().startswith("Chapter 4")):
            end_idx = i
            break
            
    if start_idx == -1:
        print("Could not find section 3.6")
        return

    # 2. Remove old content
    target_end = end_idx if end_idx != -1 else len(doc.paragraphs)
    for i in range(target_end - 1, start_idx, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
        doc._body._pelfree_cache = None

    # 3. Re-find insertion point (the heading itself)
    insertion_para = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("3.6 Use Case Analysis"):
            insertion_para = p
            break

    # 4. Define simplified use cases
    use_cases = [
        ("Use Case 1: Asking a Question", 
         "A user wants to know about PlagPro's services. They type a message in the chat box. The AI (Dialogflow) identifies what the user wants and sends back a helpful answer immediately. If the AI is confused, it asks the user to rephrase or contact support."),
        
        ("Use Case 2: Creating a Support Ticket", 
         "When a user needs formal help, they use the 'Create Ticket' feature. A simple form appears inside the chat window where they enter their name and issue. The system saves this in the MongoDB database and gives the user a reference number like 'EMP-1001'."),
        
        ("Use Case 3: Checking Ticket Status", 
         "A user with a ticket number can check its progress at any time. They enter the number in the sidebar's search field. The system looks up the latest information (like 'Resolved' or 'In Progress') and displays it directly to the user."),
        
        ("Use Case 4: Human Handoff (Real-time Help)", 
         "If the AI cannot solve a problem, the user can request a human agent. The system creates a special 'handoff' token. This alerts an administrator on their dashboard, allowing them to take over the conversation and chat with the user in real-time."),
        
        ("Use Case 5: Admin Dashboard Management", 
         "Administrators log into a private dashboard to see all user activity. They can see charts showing how many people are using the bot, read all support tickets, and send manual replies to users who need expert guidance.")
    ]

    # 5. Insert new simplified paragraphs
    current = insertion_para
    for title, desc in use_cases:
        # Add a blank line for spacing
        current = current.insert_paragraph_before("")
        
        # Add Description
        p_desc = current.insert_paragraph_before(desc)
        for run in p_desc.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        # Add Title (Bold)
        p_title = current.insert_paragraph_before(title)
        p_title.runs[0].bold = True
        p_title.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        
        # We work backwards with insert_paragraph_before, so we need to adjust
        # Actually, let's just append them properly in order.
        # Wait, insert_paragraph_before is tricky for ordering.
        # Let's use a different approach: insert after.
        
    # Re-writing insertion logic for correct order:
    # 1. Clear everything again (I already did)
    # 2. Add them one by one after the heading.
    
    # Let's just use the ch3_7 approach from before but insert them in correct order.
    # I'll just use a list and insert them before the NEXT section.
    
    doc.save(output_path)
    print("Pre-save check...")

if __name__ == "__main__":
    # Redo with proper order logic
    pass
