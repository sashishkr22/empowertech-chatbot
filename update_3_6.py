import docx
from docx.enum.text import WD_COLOR_INDEX

def update_3_6(file_path, output_path):
    doc = docx.Document(file_path)
    
    start_idx = -1
    end_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("3.6 Use Case Analysis"):
            start_idx = i
        elif start_idx != -1 and (p.text.strip().startswith("3.7") or p.text.strip().startswith("Chapter 4")):
            end_idx = i
            break
            
    if start_idx == -1:
        print("Could not find section 3.6")
        return

    # Delete existing paragraphs between 3.6 and 3.7
    # Note: end_idx might be None if 3.6 is at the end of Ch3
    target_end = end_idx if end_idx != -1 else len(doc.paragraphs)
    
    for i in range(target_end - 1, start_idx, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
        doc._body._pelfree_cache = None

    # Re-find insertion point
    ch3_7 = None
    if end_idx != -1:
        for p in doc.paragraphs:
            if p.text.strip().startswith("3.7") or p.text.strip().startswith("Chapter 4"):
                ch3_7 = p
                break

    use_cases = [
        ("Use Case 1: User Queries a Service", "Actor: Prospective client. Trigger: User types a service-related question. Normal flow: User types their message → chat.js sends POST /api/chat → Dialogflow classifies intent → Node.js returns fulfilment text → chat.js renders bot reply. Alternative flow: Confidence below threshold → Default Fallback Intent triggered → User redirected to contact options."),
        ("Use Case 2: User Creates a Support Ticket", "Actor: Client needing assistance. Trigger: User clicks 'Create Ticket' or intent detected. Normal flow: User fills redesigned inline form (Name, Email, Query) → chat.js sends POST /api/ticket/create → Node.js generates EMP-XXXX ID → Data written to MongoDB Atlas → Success message with Ticket ID displayed to user."),
        ("Use Case 3: User Checks Ticket Status", "Actor: Client with existing ticket. Trigger: User enters ID in Sidebar search. Normal flow: User types EMP-XXXX → chat.js calls GET /api/ticket/:id → Backend retrieves status and priority from MongoDB → Results rendered directly in the Sidebar status area."),
        ("Use Case 4: Human Handoff (Real-time Consultation)", "Actor: User and Support Agent. Trigger: HumanHandoff intent detected. Normal flow: Bot creates a 'handoff' token in MongoDB → Admin Dashboard alerts staff via real-time polling → Admin opens handoff detail page → Bidirectional chat synchronization enabled via periodic fetching (poll method) → Admin replies rendered in user chat interface."),
        ("Use Case 5: Admin Dashboard Analytics", "Actor: Support Manager. Trigger: Admin logs into dashboard. Normal flow: Flask app fetches aggregate ticket data from MongoDB → Dashboard renders KPI cards (Total, Open, In Progress) → Bar charts and service breakdown graphs generated dynamically using Jinja2 templates.")
    ]

    # Insert before 3.7 (or append if not found)
    current_point = ch3_7 if ch3_7 else doc.paragraphs[-1]
    
    for title, content in reversed(use_cases):
        # Spacing
        current_point.insert_paragraph_before("")
        # Content
        p_content = current_point.insert_paragraph_before(content)
        for run in p_content.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        # Title
        p_title = current_point.insert_paragraph_before(title)
        p_title.runs[0].bold = True
        p_title.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.save(output_path)
    print(f"Section 3.6 updated and saved to {output_path}")

if __name__ == "__main__":
    update_3_6("Project Report (Updated v3).docx", "Project Report (Updated v4).docx")
